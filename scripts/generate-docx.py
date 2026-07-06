import sys, re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def set_font(run, name='微软雅黑', size=Pt(10.5), bold=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_formatted_paragraph(doc, text, bold=False, size=None, alignment=None, space_before=0, space_after=4):
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    set_font(run, bold=bold)
    if size:
        run.font.size = size
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    
    # Heading styles
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '微软雅黑'
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12), 4: Pt(11)}
        hs.font.size = sizes.get(level, Pt(11))
        hs.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        hs.paragraph_format.space_after = Pt(6)
        # Set east-asia font
        rPr = hs.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Read lines
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    lines = text.split('\n')
    
    i = 0
    in_yaml = False
    in_code = False
    
    while i < len(lines):
        line = lines[i]
        
        # YAML frontmatter
        if i == 0 and line.strip() == '---':
            in_yaml = True
            i += 1
            continue
        if in_yaml:
            if line.strip() == '---':
                in_yaml = False
            i += 1
            continue
        
        # Code block
        if line.strip().startswith('```'):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue
        
        stripped = line.strip()
        
        # Headings
        if stripped.startswith('#') and stripped[1:2] in (' ', '#'):
            level = 1
            while level < 5 and stripped.startswith('#' * level + ' '):
                level += 1
            if level > 1:
                level -= 1
            heading_text = stripped.lstrip('#').strip()
            if heading_text:
                doc.add_heading(heading_text, level=level)
        
        # Table
        elif stripped.startswith('|') and stripped.endswith('|'):
            rows_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows_data.append(lines[i].strip())
                i += 1
            i -= 1  # compensate for loop increment
            
            if len(rows_data) >= 2:
                # Parse header
                headers = [h.strip() for h in rows_data[0].split('|') if h.strip()]
                # Find separator line (|----|)
                sep_idx = 1
                while sep_idx < len(rows_data) and re.match(r'^\|[-|\s]+\|$', rows_data[sep_idx]):
                    sep_idx += 1
                data_rows = rows_data[sep_idx:] if sep_idx < len(rows_data) else []
                
                if headers:
                    ncols = len(headers)
                    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Style table
                    tbl = table._tbl
                    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
                    tblStyle = tblPr.find(qn('w:tblStyle'))
                    if tblStyle is None:
                        tblStyle = OxmlElement('w:tblStyle')
                        tblPr.append(tblStyle)
                    tblStyle.set(qn('w:val'), 'Table Grid')
                    
                    # Header row
                    for j, h in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(h)
                        set_font(run, bold=True, size=Pt(9.5))
                        set_cell_shading(cell, '2B579A')
                        # White text for header
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    
                    # Data rows
                    for ri, row_text in enumerate(data_rows):
                        cells_text = [c.strip() for c in row_text.split('|') if c.strip()]
                        for j, ct in enumerate(cells_text[:ncols]):
                            cell = table.rows[ri+1].cells[j]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            # Handle **bold** markers in cell text
                            clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', ct)
                            run = p.add_run(clean_text)
                            set_font(run, size=Pt(9.5))
                            # Alternate row shading
                            if ri % 2 == 1:
                                set_cell_shading(cell, 'F2F6FC')
                    
                    # Add spacing after table
                    doc.add_paragraph().paragraph_format.space_after = Pt(2)
        
        # Horizontal rule
        elif stripped in ('---', '***', '___'):
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
        # Empty line
        elif not stripped:
            pass
        
        # Regular content
        else:
            # Handle bold markers **text**
            if '**' in stripped:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.+?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        set_font(run, bold=True)
                    else:
                        run = p.add_run(part)
                        set_font(run)
            else:
                p = doc.add_paragraph(stripped)
                for run in p.runs:
                    set_font(run)
        
        i += 1
    
    # Set default font for entire document
    doc.save(docx_path)
    return True

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python generate-docx.py <输入.md> [输出.docx]")
        sys.exit(1)
    
    md_path = sys.argv[1]
    if len(sys.argv) >= 3:
        docx_path = sys.argv[2]
    else:
        docx_path = os.path.splitext(md_path)[0] + '.docx'
    
    print(f"读取: {md_path}")
    if convert_md_to_docx(md_path, docx_path):
        print(f"已生成: {docx_path}")
    else:
        print("转换失败")
        sys.exit(1)
